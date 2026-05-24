# agent/brain.py — Cerveau de Tima : connexion avec l'API Claude + tool calling

import os
import yaml
import logging
from pathlib import Path
from anthropic import AsyncAnthropic
from dotenv import load_dotenv

load_dotenv(override=True)
logger = logging.getLogger("agentkit")

client = AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# ── Modèle actif (basculé dynamiquement par le scheduler) ─────────────────────
# Par défaut : Haiku 4.5 (économique)
# À 2h AM → Sonnet 4.6 (mises à jour SAMANTAN)
# À 2h30 AM → retour Haiku 4.5
_MODELE_PAR_DEFAUT = "claude-haiku-4-5"


def get_modele_actif() -> str:
    """Retourne le modèle Claude actuellement actif pour Tima."""
    return os.getenv("TIMA_MODEL", _MODELE_PAR_DEFAUT)


def set_modele(modele: str) -> None:
    """Change le modèle actif de Tima (appelé par le scheduler)."""
    os.environ["TIMA_MODEL"] = modele
    logger.info(f"Modèle Tima basculé → {modele}")

# ── Outil : accès au catalogue SAMANTAN ───────────────────────────────────────
TOOLS = [
    {
        "name": "consulter_catalogue_samantan",
        "description": (
            "Accède au catalogue SAMANTAN en temps réel pour vérifier les produits actifs, "
            "références, gammes et disponibilités. Utiliser si le client pose une question "
            "très spécifique sur un produit ou si les infos du contexte ne suffisent pas."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "recherche": {
                    "type": "string",
                    "description": "Type de produit recherché (ex: 'progressif', 'transitions', 'unifocal')"
                }
            },
            "required": []
        }
    },
    {
        "name": "consulter_prix_opticien",
        "description": (
            "Consulte les prix personnalisés d'un opticien spécifique : tarifs par produit, "
            "paliers de remise mensuels, paramétrage 2ème/3ème paire, pactoïe. "
            "Utiliser quand un opticien demande ses prix, ses remises ou ses avantages."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "nom_opticien": {
                    "type": "string",
                    "description": "Nom de l'opticien ou de la boutique (ex: 'OPTIQUE PONTY', 'JUNIOR OPTIQUE')"
                }
            },
            "required": []
        }
    },
    {
        "name": "consulter_ordonnances_samantan",
        "description": (
            "Accède à la liste des ordonnances/commandes du site SAMANTAN en temps réel. "
            "Utiliser quand un client demande l'état d'une commande, le suivi d'une "
            "ordonnance, des informations sur les prescriptions en cours, ou l'historique "
            "des commandes."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "recherche": {
                    "type": "string",
                    "description": "Référence de commande, nom du client, statut ou autre terme de recherche"
                }
            },
            "required": []
        }
    },
    {
        "name": "enregistrer_nouvel_opticien",
        "description": (
            "Enregistre un NOUVEL opticien (pas encore client/activé chez SAMANTAN) et "
            "transmet ses informations à l'équipe SAMANTAN pour activation de son compte. "
            "⚠️ N'appeler QU'UNE SEULE FOIS, et UNIQUEMENT quand TOUTES les infos obligatoires "
            "ont été collectées auprès de l'opticien : nom de la boutique, nom du responsable, "
            "téléphone et email. Ne JAMAIS appeler s'il manque une de ces infos obligatoires, "
            "et ne JAMAIS inventer une valeur manquante."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "nom_boutique": {
                    "type": "string",
                    "description": "Nom de la boutique ou du cabinet optique (obligatoire)"
                },
                "nom_responsable": {
                    "type": "string",
                    "description": "Nom du responsable / gérant de la boutique (obligatoire)"
                },
                "telephone": {
                    "type": "string",
                    "description": "Numéro de téléphone de contact de l'opticien (obligatoire)"
                },
                "email": {
                    "type": "string",
                    "description": "Email principal de l'opticien (obligatoire)"
                },
                "email_facturation": {
                    "type": "string",
                    "description": "Email de facturation s'il est différent de l'email principal (facultatif)"
                },
                "adresse": {
                    "type": "string",
                    "description": "Adresse ou ville de la boutique (facultatif)"
                }
            },
            "required": ["nom_boutique", "nom_responsable", "telephone", "email"]
        }
    },
    {
        "name": "transmettre_client_final",
        "description": (
            "Transmet la demande d'un CLIENT FINAL / porteur de verre (un particulier, PAS un "
            "opticien professionnel) à l'opticien partenaire YELETA OPTIC et à l'équipe SAMANTAN. "
            "SAMANTAN ne vend qu'aux opticiens : le client final doit passer par un partenaire. "
            "Appeler quand tu as collecté au minimum le téléphone et le besoin du client. "
            "Inclure aussi l'ordonnance, l'assurance, la monture souhaitée et ta recommandation "
            "si tu les as. N'invente JAMAIS une information non fournie par le client."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "nom": {
                    "type": "string",
                    "description": "Nom du client final (facultatif)"
                },
                "telephone": {
                    "type": "string",
                    "description": "Numéro de téléphone du client (obligatoire)"
                },
                "besoin": {
                    "type": "string",
                    "description": "Ce que cherche le client (ex: vision de loin, progressif, solaire) (obligatoire)"
                },
                "ordonnance": {
                    "type": "string",
                    "description": "Valeurs de l'ordonnance si fournies (OD/OG Sph/Cyl/Axe, Add...) ou 'photo reçue'"
                },
                "assurance": {
                    "type": "string",
                    "description": "Assureur / mutuelle du client s'il en a une"
                },
                "monture": {
                    "type": "string",
                    "description": "Préférence de monture / choix exprimé par le client"
                },
                "recommandation": {
                    "type": "string",
                    "description": "Conseil de Tima : produit/indice/traitement adaptés au besoin du client"
                }
            },
            "required": ["telephone", "besoin"]
        }
    }
]

# Numéro WhatsApp de l'équipe SAMANTAN qui reçoit les demandes d'inscription
# des nouveaux opticiens (overridable via la variable d'environnement ONBOARDING_PHONE).
ONBOARDING_PHONE = os.getenv("ONBOARDING_PHONE", "221775434816")

# Opticien partenaire qui reçoit les leads de clients finaux / porteurs de verre.
YELETA_PHONE = os.getenv("YELETA_PHONE", "221771961316")
YELETA_MAPS = "https://www.google.com/maps/search/?api=1&query=YELETA+OPTIC+Dakar"


def cargar_config_prompts() -> dict:
    """Lit la configuration depuis config/prompts.yaml."""
    try:
        with open("config/prompts.yaml", "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except FileNotFoundError:
        logger.error("config/prompts.yaml introuvable")
        return {}


def cargar_knowledge() -> str:
    """
    Charge les fichiers de knowledge/ dans le system prompt de Tima.

    Règles de taille pour garder le contexte Claude sous contrôle :
      - prix_opticiens.md  → EXCLU du system prompt (trop grand, ~125 KB)
                             accessible uniquement via l'outil consulter_prix_opticien
      - catalogue_samantan.md → limité à 12 000 chars (résumé des produits actifs)
      - Autres fichiers    → limités à 8 000 chars chacun
      - Total knowledge    → plafonné à 30 000 chars
    """
    knowledge_dir = Path("knowledge")
    contenus = []
    total_chars = 0
    LIMITE_TOTALE = 30_000

    # Fichiers exclus du system prompt :
    #  - prix_opticiens.md   : trop grand, accessible via l'outil consulter_prix_opticien
    #  - reseau_opticiens.md : liste des opticiens = CLIENTS de SAMANTAN, strictement
    #                          confidentielle. Ne JAMAIS l'injecter dans le contexte sinon
    #                          Tima risque de divulguer la liste des clients à un tiers.
    EXCLUS = {"prix_opticiens.md", "reseau_opticiens.md"}

    # Limites par fichier
    LIMITES = {
        "catalogue_samantan.md": 12_000,
    }
    LIMITE_DEFAUT = 8_000

    if not knowledge_dir.exists():
        return ""

    for fichier in sorted(knowledge_dir.glob("*.md")):
        if fichier.name.startswith(".") or fichier.name in EXCLUS:
            if fichier.name in EXCLUS:
                logger.info(f"Knowledge exclu (outil dédié) : {fichier.name}")
            continue

        if total_chars >= LIMITE_TOTALE:
            logger.info(f"Limite knowledge atteinte ({LIMITE_TOTALE} chars) — {fichier.name} ignoré")
            break

        try:
            texte = fichier.read_text(encoding="utf-8").strip()
            if not texte or len(texte) < 50:
                continue

            limite = LIMITES.get(fichier.name, LIMITE_DEFAUT)
            if len(texte) > limite:
                texte = texte[:limite] + f"\n\n_(... tronqué à {limite} chars)_"

            contenus.append(f"\n\n---\n{texte}")
            total_chars += len(texte)
            logger.info(f"Knowledge chargé : {fichier.name} ({len(texte)} chars)")
        except Exception as e:
            logger.warning(f"Impossible de lire {fichier} : {e}")

    logger.info(f"Knowledge total : {total_chars} chars / {LIMITE_TOTALE} max")
    return "\n".join(contenus)


# ── Cache du system prompt en mémoire (recalculé uniquement si les fichiers changent) ──
_system_prompt_cache: str | None = None


def cargar_system_prompt() -> str:
    global _system_prompt_cache
    if _system_prompt_cache is not None:
        return _system_prompt_cache

    config = cargar_config_prompts()
    system_prompt = config.get(
        "system_prompt",
        "Tu es Tima, assistante de SAMANTAN. Réponds en français."
    )

    # Enrichir avec le contenu du site SAMANTAN (knowledge/)
    knowledge = cargar_knowledge()
    if knowledge:
        system_prompt += (
            "\n\n## Contenu récupéré depuis le site SAMANTAN\n"
            "Utilise ces informations pour répondre aux questions sur SAMANTAN.\n"
            f"{knowledge}"
        )

    _system_prompt_cache = system_prompt
    logger.info(f"System prompt mis en cache ({len(system_prompt)} chars)")
    return system_prompt


def invalider_cache_system_prompt():
    """Invalide le cache du system prompt (appeler après une mise à jour des knowledge files)."""
    global _system_prompt_cache
    _system_prompt_cache = None
    logger.info("Cache system prompt invalidé")


def obtener_mensaje_error() -> str:
    config = cargar_config_prompts()
    return config.get("error_message", "Désolée, problème technique. Veuillez réessayer.")


def obtener_mensaje_fallback() -> str:
    config = cargar_config_prompts()
    return config.get("fallback_message", "Désolée, je n'ai pas bien compris. Tu peux reformuler ?")


def _extraire_texte(content: list) -> str:
    """
    Extrait le texte du premier bloc TextBlock dans une liste de blocs Anthropic.
    Sûr même si le premier bloc est un ToolUseBlock ou si la liste est vide.
    """
    for bloc in content:
        if hasattr(bloc, "type") and bloc.type == "text" and hasattr(bloc, "text"):
            return bloc.text
    return ""


async def _executer_outil(nom: str, parametres: dict, telefono: str = "") -> str:
    """Exécute un outil demandé par Claude et retourne le résultat."""

    if nom == "consulter_catalogue_samantan":
        from agent.web_scraper import fetch_catalogue_samantan
        recherche = parametres.get("recherche", "")
        logger.info(f"Tima consulte le catalogue SAMANTAN (recherche: '{recherche}')")
        return await fetch_catalogue_samantan(recherche)

    if nom == "consulter_prix_opticien":
        nom_opticien = parametres.get("nom_opticien", "").strip()
        logger.info(f"Tima consulte les prix de : '{nom_opticien}'")
        return _chercher_prix_opticien(nom_opticien)

    if nom == "consulter_ordonnances_samantan":
        from agent.web_scraper import fetch_ordonnances_samantan
        recherche = parametres.get("recherche", "")
        logger.info(f"Tima consulte les ordonnances SAMANTAN (recherche: '{recherche}')")
        return await fetch_ordonnances_samantan(recherche)

    if nom == "enregistrer_nouvel_opticien":
        return await _enregistrer_nouvel_opticien(parametres, telefono)

    if nom == "transmettre_client_final":
        return await _transmettre_client_final(parametres, telefono)

    return "Outil inconnu."


def _enregistrer_demande_localement(infos: dict, fichier: str = "knowledge/demandes_inscription.json") -> None:
    """
    Sauvegarde une demande (inscription opticien ou lead client final) dans un
    fichier JSON local — file d'attente utile si l'envoi WhatsApp échoue ou pour archive.
    """
    import json
    from datetime import datetime
    try:
        chemin = Path(fichier)
        chemin.parent.mkdir(parents=True, exist_ok=True)
        demandes = []
        if chemin.exists():
            try:
                demandes = json.loads(chemin.read_text(encoding="utf-8"))
            except Exception:
                demandes = []
        infos = {**infos, "date": datetime.now().strftime("%Y-%m-%d %H:%M")}
        demandes.append(infos)
        chemin.write_text(json.dumps(demandes, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        logger.warning(f"Impossible d'enregistrer la demande localement : {e}")


async def _enregistrer_nouvel_opticien(params: dict, telefono_expediteur: str = "") -> str:
    """
    Transmet la demande d'inscription d'un nouvel opticien à l'équipe SAMANTAN
    (numéro ONBOARDING_PHONE) via WhatsApp, et archive la demande localement.
    """
    nom_boutique     = (params.get("nom_boutique") or "").strip()
    nom_responsable  = (params.get("nom_responsable") or "").strip()
    telephone        = (params.get("telephone") or "").strip() or telefono_expediteur
    email            = (params.get("email") or "").strip()
    email_facturation = (params.get("email_facturation") or "").strip()
    adresse          = (params.get("adresse") or "").strip()

    # Sécurité : vérifier que les infos obligatoires sont bien présentes
    manquants = [n for n, v in [
        ("nom de la boutique", nom_boutique),
        ("nom du responsable", nom_responsable),
        ("téléphone", telephone),
        ("email", email),
    ] if not v]
    if manquants:
        return f"INFOS_MANQUANTES: Il manque encore : {', '.join(manquants)}. Demande-les avant d'enregistrer."

    # ── Construire le message destiné à l'équipe SAMANTAN ──────────────────────
    lignes = [
        "🆕 NOUVELLE DEMANDE D'INSCRIPTION OPTICIEN",
        "",
        f"Boutique     : {nom_boutique}",
        f"Responsable  : {nom_responsable}",
        f"Téléphone    : {telephone}",
        f"Email        : {email}",
    ]
    if email_facturation:
        lignes.append(f"Email factu. : {email_facturation}")
    if adresse:
        lignes.append(f"Adresse      : {adresse}")
    if telefono_expediteur:
        lignes.append(f"WhatsApp     : {telefono_expediteur}")
    lignes.append("")
    lignes.append("→ Compte à activer dans l'espace SAMANTAN.")
    message_admin = "\n".join(lignes)

    infos = {
        "nom_boutique": nom_boutique,
        "nom_responsable": nom_responsable,
        "telephone": telephone,
        "email": email,
        "email_facturation": email_facturation,
        "adresse": adresse,
        "whatsapp": telefono_expediteur,
    }
    _enregistrer_demande_localement(infos)

    # ── Envoyer à l'équipe SAMANTAN via le proveedor WhatsApp ──────────────────
    try:
        from agent.providers import obtener_proveedor
        proveedor = obtener_proveedor()
        envoye = await proveedor.enviar_mensaje(ONBOARDING_PHONE, message_admin)
        if envoye:
            logger.info(f"Demande d'inscription transmise à l'équipe SAMANTAN ({ONBOARDING_PHONE}) ✓")
            return (
                "INSCRIPTION_TRANSMISE: Les informations ont bien été transmises à l'équipe "
                "SAMANTAN. Dis maintenant à l'opticien d'attendre l'activation de son compte — "
                "un opérateur reviendra vers lui."
            )
        logger.error(f"Échec envoi demande d'inscription à {ONBOARDING_PHONE}")
        return (
            "INSCRIPTION_NOTEE: Les informations sont bien notées et seront transmises à "
            "l'équipe SAMANTAN. Dis à l'opticien d'attendre l'activation de son compte."
        )
    except Exception as e:
        logger.error(f"Erreur transmission demande d'inscription : {e}")
        return (
            "INSCRIPTION_NOTEE: Les informations sont bien notées et seront transmises à "
            "l'équipe SAMANTAN. Dis à l'opticien d'attendre l'activation de son compte."
        )


async def _transmettre_client_final(params: dict, telefono_expediteur: str = "") -> str:
    """
    Transmet le lead d'un client final / porteur de verre à l'opticien partenaire
    YELETA OPTIC (YELETA_PHONE) ET à l'équipe SAMANTAN (ONBOARDING_PHONE), via WhatsApp.
    Archive également la demande localement.
    """
    nom            = (params.get("nom") or "").strip()
    telephone      = (params.get("telephone") or "").strip() or telefono_expediteur
    besoin         = (params.get("besoin") or "").strip()
    ordonnance     = (params.get("ordonnance") or "").strip()
    assurance      = (params.get("assurance") or "").strip()
    monture        = (params.get("monture") or "").strip()
    recommandation = (params.get("recommandation") or "").strip()

    # Sécurité : téléphone + besoin obligatoires
    manquants = [n for n, v in [("téléphone", telephone), ("besoin", besoin)] if not v]
    if manquants:
        return f"INFOS_MANQUANTES: Il manque encore : {', '.join(manquants)}. Demande-les avant de transmettre."

    # ── Construire le message pour YELETA OPTIC + équipe SAMANTAN ──────────────
    lignes = ["👓 NOUVEAU CLIENT FINAL (porteur de verre)", ""]
    if nom:
        lignes.append(f"Nom        : {nom}")
    lignes.append(f"Téléphone  : {telephone}")
    lignes.append(f"Besoin     : {besoin}")
    if ordonnance:
        lignes.append(f"Ordonnance : {ordonnance}")
    if assurance:
        lignes.append(f"Assurance  : {assurance}")
    if monture:
        lignes.append(f"Monture    : {monture}")
    if recommandation:
        lignes.append(f"Conseil Tima : {recommandation}")
    if telefono_expediteur:
        lignes.append(f"WhatsApp   : {telefono_expediteur}")
    lignes.append("")
    lignes.append("→ Client à recontacter pour finaliser sa commande de verres.")
    message = "\n".join(lignes)

    _enregistrer_demande_localement({
        "nom": nom, "telephone": telephone, "besoin": besoin,
        "ordonnance": ordonnance, "assurance": assurance, "monture": monture,
        "recommandation": recommandation, "whatsapp": telefono_expediteur,
    }, fichier="knowledge/demandes_clients_finaux.json")

    # ── Envoyer à YELETA OPTIC + équipe SAMANTAN ───────────────────────────────
    try:
        from agent.providers import obtener_proveedor
        proveedor = obtener_proveedor()
        resultats = []
        for dest in (YELETA_PHONE, ONBOARDING_PHONE):
            try:
                ok = await proveedor.enviar_mensaje(dest, message)
            except Exception as e:
                logger.error(f"Erreur envoi lead client final à {dest} : {e}")
                ok = False
            resultats.append(ok)
            logger.info(f"Lead client final → {dest} : {'OK' if ok else 'ÉCHEC'}")

        if any(resultats):
            return (
                "CLIENT_TRANSMIS: La demande a été transmise à l'opticien partenaire YELETA OPTIC "
                "et à l'équipe SAMANTAN. Dis au client que YELETA OPTIC le recontactera pour "
                "finaliser, et qu'il peut aussi s'y rendre directement."
            )
        return (
            "CLIENT_NOTE: La demande est bien notée et sera transmise à YELETA OPTIC et à l'équipe "
            "SAMANTAN. Dis au client qu'il sera recontacté, et oriente-le vers YELETA OPTIC."
        )
    except Exception as e:
        logger.error(f"Erreur transmission client final : {e}")
        return (
            "CLIENT_NOTE: La demande est bien notée et sera transmise à YELETA OPTIC et à l'équipe "
            "SAMANTAN. Dis au client qu'il sera recontacté, et oriente-le vers YELETA OPTIC."
        )


def _chercher_prix_opticien(nom_opticien: str) -> str:
    """
    Recherche les prix d'un opticien dans knowledge/prix_opticiens.md.
    Retourne la section correspondante ou les 5 premiers opticiens si nom vide.
    """
    try:
        fichier = Path("knowledge/prix_opticiens.md")
        if not fichier.exists():
            return "Fichier de prix non disponible — relance /scraper-prix-opticiens."

        contenu = fichier.read_text(encoding="utf-8")
        sections = contenu.split("\n## ")

        if not nom_opticien:
            # Retourner les 3 premiers opticiens comme exemple
            exemples = sections[1:4] if len(sections) > 1 else []
            return "Exemples de prix :\n\n## " + "\n\n## ".join(exemples) if exemples else contenu[:3000]

        # Recherche partielle insensible à la casse
        nom_lower = nom_opticien.lower()
        for section in sections[1:]:
            titre = section.split("\n")[0].lower()
            if nom_lower in titre:
                return "## " + section[:4000]

        # Pas trouvé → proposer les noms disponibles
        noms = [s.split("\n")[0] for s in sections[1:] if s.strip()][:20]
        return (
            f"Opticien '{nom_opticien}' non trouvé.\n"
            f"Opticiens disponibles : {', '.join(noms)}"
        )
    except Exception as e:
        logger.warning(f"Erreur lecture prix opticiens : {e}")
        return "Impossible de lire les prix pour l'instant."


async def _telecharger_image_base64(url: str) -> tuple[str, str] | None:
    """
    Télécharge une image depuis une URL et retourne (base64, media_type).
    Retourne None si échec.
    """
    import base64
    import httpx
    try:
        async with httpx.AsyncClient(timeout=20.0) as http:
            r = await http.get(url)
            if r.status_code != 200:
                logger.warning(f"Image inaccessible : HTTP {r.status_code}")
                return None
            content_type = r.headers.get("content-type", "image/jpeg")
            # Normaliser le media_type pour Claude
            if "png" in content_type:
                media_type = "image/png"
            elif "webp" in content_type:
                media_type = "image/webp"
            elif "gif" in content_type:
                media_type = "image/gif"
            else:
                media_type = "image/jpeg"
            b64 = base64.standard_b64encode(r.content).decode("utf-8")
            logger.info(f"Image téléchargée : {len(r.content)} bytes ({media_type})")
            return b64, media_type
    except Exception as e:
        logger.warning(f"Erreur téléchargement image : {e}")
        return None


async def _telecharger_fichier(url: str) -> bytes | None:
    """Télécharge un fichier depuis une URL et retourne ses bytes bruts. None si échec."""
    import httpx
    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as http:
            r = await http.get(url)
            if r.status_code != 200:
                logger.warning(f"Fichier inaccessible : HTTP {r.status_code}")
                return None
            logger.info(f"Fichier téléchargé : {len(r.content)} bytes")
            return r.content
    except Exception as e:
        logger.warning(f"Erreur téléchargement fichier : {e}")
        return None


def _extraire_texte_docx(contenu: bytes) -> str:
    """Extrait le texte d'un fichier Word (.docx)."""
    import io
    try:
        from docx import Document
        doc = Document(io.BytesIO(contenu))
        lignes = [p.text for p in doc.paragraphs if p.text.strip()]
        # Extraire aussi les tableaux
        for table in doc.tables:
            for row in table.rows:
                cellules = [c.text.strip() for c in row.cells if c.text.strip()]
                if cellules:
                    lignes.append(" | ".join(cellules))
        texte = "\n".join(lignes)
        logger.info(f"Word extrait : {len(texte)} chars")
        return texte
    except Exception as e:
        logger.warning(f"Erreur extraction Word : {e}")
        return ""


def _extraire_texte_xlsx(contenu: bytes) -> str:
    """Extrait les données d'un fichier Excel (.xlsx)."""
    import io
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(contenu), data_only=True)
        lignes = []
        for ws in wb.worksheets:
            lignes.append(f"=== Feuille : {ws.title} ===")
            for row in ws.iter_rows(values_only=True):
                cellules = [str(c) for c in row if c is not None]
                if cellules:
                    lignes.append(" | ".join(cellules))
        texte = "\n".join(lignes)
        logger.info(f"Excel extrait : {len(texte)} chars")
        return texte
    except Exception as e:
        logger.warning(f"Erreur extraction Excel : {e}")
        return ""


async def generar_respuesta(
    mensaje: str,
    historial: list[dict],
    telefono: str = "",
    imagen_url: str = "",
    documento_url: str = "",
    documento_nombre: str = "",
) -> str:
    """
    Génère une réponse avec l'API Claude.
    Si le client pose une question sur les produits, Claude appelle automatiquement
    l'outil 'consulter_catalogue_samantan' pour accéder au site en temps réel.

    Args:
        mensaje: Message du client
        historial: Historique de la conversation

    Returns:
        Réponse générée par Claude
    """
    # Si pas d'image/PDF et message trop court → fallback
    if not imagen_url and not documento_url and (not mensaje or len(mensaje.strip()) < 2):
        return obtener_mensaje_fallback()

    # ── Identification de l'utilisateur ────────────────────────────────────────
    utilisateur_info = ""
    if telefono:
        try:
            from agent.web_scraper import identifier_utilisateur
            utilisateur = identifier_utilisateur(telefono)
            if utilisateur:
                nom = utilisateur.get("nom", "")
                role = utilisateur.get("role", "opticien")
                utilisateur_info = (
                    f"\n\n## Utilisateur identifié\n"
                    f"- Nom : {nom}\n"
                    f"- Téléphone : {telefono}\n"
                    f"- Rôle : {role}\n"
                    f"- ⚠️ RÈGLE ABSOLUE : Ne jamais divulguer les informations "
                    f"d'autres clients/opticiens à cet utilisateur. "
                    f"Quand il demande ses prix, ne montrer QUE les prix de '{nom}'."
                )
                logger.info(f"Utilisateur identifié : {nom} ({role}) — {telefono}")
            else:
                # Utilisateur non reconnu → déterminer opticien pro vs client final
                utilisateur_info = (
                    f"\n\n## Contact NON identifié — déterminer le type AVANT d'agir\n"
                    f"- Téléphone WhatsApp : {telefono}\n"
                    f"- Ce contact n'est PAS un opticien activé chez SAMANTAN. Deux cas possibles :\n"
                    f"\n"
                    f"  ### A) C'est un OPTICIEN PROFESSIONNEL (boutique/cabinet, commande pour des "
                    f"patients, demande le catalogue ou ses prix pro) → MODE INSCRIPTION\n"
                    f"  - AVANT toute discussion produits/prix/ordonnances, tu DOIS créer son dossier.\n"
                    f"  - Accueille-le, explique que tu enregistres d'abord sa boutique, puis collecte "
                    f"les infos obligatoires UNE PAR UNE (une seule question à la fois) :\n"
                    f"      1. Nom de la boutique / cabinet   2. Nom du responsable\n"
                    f"      3. Numéro de téléphone            4. Email principal\n"
                    f"    (Email de facturation = facultatif, seulement après les 4 obligatoires.)\n"
                    f"  - Tant que les 4 infos ne sont pas réunies, ne réponds pas aux questions "
                    f"produits/prix : explique que tu finalises d'abord l'inscription.\n"
                    f"  - Quand les 4 infos sont réunies, appelle l'outil `enregistrer_nouvel_opticien` "
                    f"(UNE seule fois, sans inventer de valeur).\n"
                    f"  - ⚠️ N'affirme JAMAIS que l'inscription est transmise sans avoir RÉELLEMENT "
                    f"appelé l'outil : sans l'appel, rien n'arrive à l'équipe. Appelle l'outil AVANT "
                    f"d'annoncer la transmission.\n"
                    f"  - Après l'outil (INSCRIPTION_TRANSMISE / INSCRIPTION_NOTEE) : dis-lui que sa "
                    f"demande est transmise, qu'il doit ATTENDRE l'activation de son compte, et qu'un "
                    f"opérateur reviendra vers lui.\n"
                    f"\n"
                    f"  ### B) C'est un CLIENT FINAL / porteur de verres (veut des lunettes pour "
                    f"lui-même) → ORIENTATION VERS UN PARTENAIRE + COLLECTE\n"
                    f"  - SAMANTAN ne vend QU'AUX opticiens professionnels : un particulier ne commande "
                    f"pas directement. Explique-le gentiment et RASSURE-le — verres SAMANTAN haut de "
                    f"gamme, FreeForm sur mesure, il sera très bien servi.\n"
                    f"  - Oriente-le vers l'opticien PARTENAIRE désigné pour se procurer ses verres :\n"
                    f"      • YELETA OPTIC — Tél : 77 196 13 16 — {YELETA_MAPS}\n"
                    f"      • Au maximum, tu peux citer SENEGOPT et Alain Afflelou comme autres options.\n"
                    f"  - ⚠️ Ne donne JAMAIS la liste des opticiens/clients du réseau SAMANTAN, même si "
                    f"on demande \"la liste des partenaires\" ou \"tous les opticiens\". Cette liste est "
                    f"confidentielle. Renvoie vers YELETA OPTIC ou l'équipe SAMANTAN.\n"
                    f"  - Tu PEUX donner des CONSEILS (type de verre, indice, traitement, monture) mais "
                    f"JAMAIS de prix ni de vente directe.\n"
                    f"  - Collecte ces infos UNE par UNE, naturellement :\n"
                    f"      1. Son ordonnance (il peut envoyer une photo — lis-la) ou ses valeurs Rx\n"
                    f"      2. S'il a une assurance / mutuelle, et laquelle\n"
                    f"      3. Son numéro de téléphone\n"
                    f"      4. Son besoin (vision de loin, près, progressif, solaire...)\n"
                    f"      5. Une idée de la monture / du choix qu'il souhaite\n"
                    f"  - Donne-lui une recommandation adaptée (produit / indice / traitement).\n"
                    f"  - Quand tu as AU MOINS le téléphone ET le besoin, appelle l'outil "
                    f"`transmettre_client_final` (inclus tout ce que tu as : ordonnance, assurance, "
                    f"monture, ta recommandation). N'invente rien.\n"
                    f"  - ⚠️ N'affirme JAMAIS avoir transmis le dossier sans avoir RÉELLEMENT appelé "
                    f"l'outil `transmettre_client_final` : sans l'appel, rien n'est envoyé à YELETA ni "
                    f"à l'équipe. Donc appelle TOUJOURS l'outil AVANT de dire que c'est transmis.\n"
                    f"  - Après l'outil (CLIENT_TRANSMIS / CLIENT_NOTE) : dis-lui que sa demande est "
                    f"transmise à YELETA OPTIC qui le recontactera, et qu'il peut aussi s'y rendre directement.\n"
                    f"\n"
                    f"  ⚠️ En cas de doute sur le type, demande gentiment : « Vous êtes opticien "
                    f"professionnel, ou vous cherchez des lunettes pour vous-même ? »\n"
                    f"  ⚠️ Ne donne AUCUN prix ni donnée confidentielle à un contact non activé."
                )
                logger.info(f"Contact non reconnu → triage opticien/client : {telefono}")
        except Exception as e:
            logger.warning(f"Erreur identification utilisateur {telefono} : {e}")

    system_prompt = cargar_system_prompt()
    if utilisateur_info:
        system_prompt = system_prompt + utilisateur_info

    mensajes = []
    for msg in historial:
        mensajes.append({"role": msg["role"], "content": msg["content"]})

    # ── Construction du message courant (document, puis image) ────────────────
    if documento_url:
        # Détecter le type via le NOM de fichier (l'URL S3 n'a pas d'extension),
        # avec fallback sur l'URL au cas où.
        url_lower = (documento_nombre + " " + documento_url).lower()
        consigne_ordonnance = (
            "Si c'est une ordonnance optique, lis attentivement et extrais les valeurs "
            "présentes : OD (Sph/Cyl/Axe), OG (Sph/Cyl/Axe), Addition, DIP. "
            "N'invente aucune valeur absente. 'pl'/'plan' = sphère 0.00. "
            "Demande confirmation avant toute commande."
        )
        contenu_fichier = await _telecharger_fichier(documento_url)

        if not contenu_fichier:
            mensajes.append({
                "role": "user",
                "content": mensaje.strip() or "J'ai reçu un fichier mais je n'arrive pas à l'ouvrir. Peux-tu le renvoyer ?",
            })
        elif url_lower.endswith(".pdf") or ".pdf" in url_lower:
            # PDF → bloc document natif Claude
            import base64
            pdf_b64 = base64.standard_b64encode(contenu_fichier).decode("utf-8")
            mensajes.append({
                "role": "user",
                "content": [
                    {"type": "document", "source": {
                        "type": "base64", "media_type": "application/pdf", "data": pdf_b64}},
                    {"type": "text", "text": mensaje.strip() or f"Voici un document PDF. {consigne_ordonnance}"},
                ],
            })
        elif url_lower.endswith((".docx", ".doc")) or ".docx" in url_lower or ".doc" in url_lower:
            # Word → extraction texte
            texte_doc = _extraire_texte_docx(contenu_fichier)
            if texte_doc:
                mensajes.append({
                    "role": "user",
                    "content": f"{mensaje.strip()}\n\n[Contenu du document Word reçu]\n{texte_doc}\n\n{consigne_ordonnance}".strip(),
                })
            else:
                mensajes.append({"role": "user", "content": "J'ai reçu un document Word mais je n'arrive pas à lire son contenu. Peux-tu me l'envoyer en PDF ou en photo ?"})
        elif url_lower.endswith((".xlsx", ".xls")) or ".xlsx" in url_lower or ".xls" in url_lower:
            # Excel → extraction données
            texte_xls = _extraire_texte_xlsx(contenu_fichier)
            if texte_xls:
                mensajes.append({
                    "role": "user",
                    "content": f"{mensaje.strip()}\n\n[Contenu du fichier Excel reçu]\n{texte_xls}\n\n{consigne_ordonnance}".strip(),
                })
            else:
                mensajes.append({"role": "user", "content": "J'ai reçu un fichier Excel mais je n'arrive pas à lire son contenu. Peux-tu me l'envoyer en PDF ou en photo ?"})
        else:
            # Type inconnu → tenter comme PDF par défaut
            mensajes.append({"role": "user", "content": mensaje.strip() or "J'ai reçu un fichier dont je ne reconnais pas le format. Tu peux me l'envoyer en PDF ou en photo ?"})
    elif imagen_url:
        img = await _telecharger_image_base64(imagen_url)
        if img:
            b64, media_type = img
            contenu_user = [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": b64,
                    },
                },
                {
                    "type": "text",
                    "text": mensaje.strip() or (
                        "Voici une photo. Si c'est une ordonnance optique, lis attentivement "
                        "et extrais les valeurs : OD (Sph/Cyl/Axe), OG (Sph/Cyl/Axe), "
                        "Addition, DIP, hauteur. Présente-les clairement et demande "
                        "confirmation avant toute commande."
                    ),
                },
            ]
            mensajes.append({"role": "user", "content": contenu_user})
        else:
            # Échec téléchargement image → traiter comme texte seul
            mensajes.append({
                "role": "user",
                "content": mensaje.strip() or "J'ai reçu une image mais je n'arrive pas à l'ouvrir. Peux-tu la renvoyer ?",
            })
    else:
        mensajes.append({"role": "user", "content": mensaje})

    try:
        # ── Appel Claude avec outils (boucle multi-tours) ──────────────────
        messages_en_cours = mensajes[:]
        MAX_TOURS = 5  # sécurité anti-boucle infinie

        # Prompt caching Anthropic : le system prompt est mis en cache côté serveur
        # → -90% sur les tokens d'entrée après le 1er appel (cache valide 5 minutes)
        system_avec_cache = [
            {
                "type": "text",
                "text": system_prompt,
                "cache_control": {"type": "ephemeral"}
            }
        ]

        modele = get_modele_actif()
        logger.info(f"Modèle actif : {modele}")

        response = await client.messages.create(
            model=modele,
            max_tokens=1024,
            system=system_avec_cache,
            messages=messages_en_cours,
            tools=TOOLS
        )

        for tour in range(MAX_TOURS):
            # ── Réponse textuelle directe → on retourne ────────────────────
            if response.stop_reason != "tool_use":
                respuesta = _extraire_texte(response.content)
                if not respuesta:
                    respuesta = obtener_mensaje_fallback()
                logger.info(
                    f"Réponse finale tour {tour} "
                    f"({response.usage.input_tokens} in / {response.usage.output_tokens} out)"
                )
                return respuesta

            # ── Claude veut utiliser un ou plusieurs outils ────────────────
            logger.info(f"Tour {tour+1} : Claude utilise un outil...")

            # Collecter TOUS les blocs tool_use de cette réponse
            tool_blocks = [b for b in response.content if b.type == "tool_use"]

            if not tool_blocks:
                break

            # Exécuter chaque outil et construire les tool_results
            tool_results = []
            for tool_block in tool_blocks:
                resultat = await _executer_outil(tool_block.name, tool_block.input, telefono)
                logger.info(
                    f"  ↳ {tool_block.name} → {len(resultat)} chars"
                )
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tool_block.id,
                    "content": resultat
                })

            # Ajouter la réponse assistant + les résultats des outils
            messages_en_cours = messages_en_cours + [
                {"role": "assistant", "content": response.content},
                {"role": "user", "content": tool_results}
            ]

            # Appel suivant avec les résultats des outils (cache system prompt réutilisé)
            response = await client.messages.create(
                model=modele,
                max_tokens=1024,
                system=system_avec_cache,
                messages=messages_en_cours,
                tools=TOOLS
            )

        # ── Sécurité : si on a épuisé les tours ───────────────────────────
        respuesta = _extraire_texte(response.content)
        if not respuesta:
            respuesta = obtener_mensaje_error()
        return respuesta

    except Exception as e:
        logger.error(f"Erreur API Claude : {type(e).__name__}: {e}")
        return obtener_mensaje_error()
